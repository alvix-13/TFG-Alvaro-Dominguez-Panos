# Manipulación de datos
import pandas as pd
import numpy as np

# Visualización de datos
import matplotlib
matplotlib.use('Agg')

# Machine Learning
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import RobustScaler, MinMaxScaler
from sklearn.compose import ColumnTransformer
from sklearn.model_selection import train_test_split, GridSearchCV
from sklearn.metrics import classification_report, accuracy_score, ConfusionMatrixDisplay, confusion_matrix
from xgboost import XGBClassifier
from sklearn.svm import SVC
from scipy.stats import randint as sp_randint

# Multiprocesado
import multiprocessing

# Guardar Modelo
import pickle

# Informe entrenamiento
from docx import Document
from datetime import datetime
import io


def cargar_datos(path):
    df = pd.read_csv(path)
    np.random.seed(33)

    X_df = df.drop(["label", "TimeStamp"], axis=1)
    y_df = df["label"]

    return X_df, y_df

def crear_clasificador(X_df, y_df):
    x_train, x_test, y_train, y_test = train_test_split(X_df, y_df, shuffle=False, test_size=0.3)

    num_cols = x_train.select_dtypes(exclude="object").columns

    num_transformer = Pipeline(steps=[
        ('scaler', MinMaxScaler()),
        ('scaler2', RobustScaler())
    ])

    preprocessor = ColumnTransformer(
        transformers=[
            ('num', num_transformer, num_cols)
    ])

    boost_pipe = Pipeline(steps=[('preprocessor', preprocessor),
                      ('classifier', XGBClassifier(random_state=33))])

    XGBoost_dist = {'classifier__max_depth'      : [3, 5, 7, 9, 11, 13, 15, 25, 50, 75, 100],
              'classifier__subsample'        : [0.1, 0.3, 0.5, 0.7, 0.9, 1],
              'classifier__learning_rate'    : [0.01, 0.1, 0.3, 0.5, 0.7, 0.9, 1],
              'classifier__n_estimators'     : sp_randint(10, 500).rvs(10)
              }

    XGBoost_grid = GridSearchCV(boost_pipe,
        param_grid= XGBoost_dist,
        scoring    = 'balanced_accuracy',
        n_jobs     = multiprocessing.cpu_count() - 1,
        cv         = 3,
        refit      = True,
        verbose    = 0
    )

    np.random.seed(33)

    XGBoost_grid.fit(X = x_train, y = y_train)

    best_params = XGBoost_grid.best_params_
    xgb_classifier = XGBClassifier(
        learning_rate=best_params['classifier__learning_rate'],
        max_depth=best_params['classifier__max_depth'],
        n_estimators=best_params['classifier__n_estimators'],
        subsample=best_params['classifier__subsample']
    )

    clasificador = Pipeline(steps=[
        ('preprocessor', preprocessor),
        ('classifier', xgb_classifier)
    ])

    pickle.dump(clasificador.fit(x_train, y_train), open('APP/modelo_XGB.pkl', 'wb'))

    return x_test, y_test

def entrenar_clasificador(clasificador, X_df, y_df):
    x_train, x_test, y_train, y_test = train_test_split(X_df, y_df, shuffle=False, test_size=0.3)

    pickle.dump(clasificador.fit(x_train, y_train), open('APP/modelo_XGB.pkl', 'wb'))

    return x_test, y_test

def ejecutar_clasificador(clasificador, x_test, y_test, doc):
    y_test_pred = clasificador.predict(x_test)

    doc.add_paragraph(classification_report(y_test, y_test_pred))

    xgb_accuracy = accuracy_score(y_test, y_test_pred)

    doc.add_paragraph(f"El valor de accuracy: {xgb_accuracy}")

    cm = confusion_matrix(y_test, y_test_pred)
    disp= ConfusionMatrixDisplay(confusion_matrix=cm, display_labels=np.unique(y_test))
    fig = disp.plot()

    img_buffer = io.BytesIO()
    fig.figure_.savefig(img_buffer, format='png')
    img_buffer.seek(0)

    doc.add_paragraph("Matriz de Confusión")

    doc.add_picture(img_buffer)

    return xgb_accuracy

def ejecutar(csv_train):
    doc = Document()

    fecha_actual = datetime.now().strftime("%d-%m-%Y")
    hora_actual = datetime.now().strftime("%H-%M-%S")
    doc.add_heading(f'Informe de entrenamiento {fecha_actual}, a las {hora_actual}', level=1)

    X_df, y_df = cargar_datos(path = csv_train)
    x_test, y_test = crear_clasificador(X_df, y_df)
    clasificador = pickle.load(open('APP/modelo_XGB.pkl','rb'))
    precision_entrenamiento = ejecutar_clasificador(clasificador, x_test, y_test, doc)

    doc.add_page_break()
    doc.add_heading(f'Informe de generalizacion {fecha_actual}, a las {hora_actual}', level=1)

    X_df, y_df = cargar_datos(path = "APP/data/data_test.csv")
    precision_generalizacion = ejecutar_clasificador(clasificador, X_df, y_df, doc)

    doc.save(f'{fecha_actual}_{hora_actual}.docx')

    return precision_entrenamiento, precision_generalizacion

def ejecutar_rapido(csv_train):
    doc = Document()

    fecha_actual = datetime.now().strftime("%d-%m-%Y")
    hora_actual = datetime.now().strftime("%H-%M-%S")
    doc.add_heading(f'Informe de entrenamiento {fecha_actual}, a las {hora_actual}', level=1)

    X_df, y_df = cargar_datos(path = csv_train)
    clasificador_viejo = pickle.load(open('APP/modelo_XGB.pkl','rb'))
    x_test, y_test = entrenar_clasificador(clasificador_viejo, X_df, y_df)

    clasificador = pickle.load(open('APP/modelo_XGB.pkl','rb'))
    precision_entrenamiento = ejecutar_clasificador(clasificador, x_test, y_test, doc)

    doc.add_page_break()
    doc.add_heading(f'Informe de generalizacion {fecha_actual}, a las {hora_actual}', level=1)

    X_df, y_df = cargar_datos(path = "APP/data/data_test.csv")
    precision_generalizacion = ejecutar_clasificador(clasificador, X_df, y_df, doc)

    doc.save(f'APP/logs/log_{fecha_actual}_{hora_actual}.docx')

    return precision_entrenamiento, precision_generalizacion